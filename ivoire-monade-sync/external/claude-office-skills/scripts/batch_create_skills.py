#!/usr/bin/env python3
"""
Batch Create/Update Claude Office Skills
Based on architecture diagram: Skills = Domain Knowledge + Templates + Scripts
"""

import os
from pathlib import Path
from datetime import datetime

# Base directory
BASE_DIR = Path(__file__).parent.parent

# Skill templates based on architecture diagram categories
SKILLS_TO_CREATE = {
    # ========== Workflow Skills (Missing) ==========
    "n8n-workflow": {
        "name": "n8n-workflow",
        "description": "Automate document workflows with n8n - 7800+ workflow templates",
        "tags": ["workflow", "automation", "n8n", "integration"],
        "library": {
            "name": "n8n",
            "url": "https://github.com/n8n-io/n8n",
            "stars": "172k"
        },
        "overview": """This skill enables document workflow automation using **n8n** - the most popular workflow automation platform with 7800+ community templates. Chain document operations, integrate with 400+ apps, and build complex document pipelines.""",
        "use_cases": [
            "Automate PDF → OCR → Translation → Email workflow",
            "Watch folder for new contracts → Review → Notify Slack",
            "Daily report generation from multiple data sources",
            "Batch document processing with conditional logic"
        ],
        "core_concepts": """
### n8n Fundamentals

n8n uses a node-based workflow approach:

```
Trigger → Action → Action → Output
   │         │         │
   └─────────┴─────────┴── Data flows between nodes
```

### Key Node Types

| Type | Examples | Use Case |
|------|----------|----------|
| **Triggers** | Webhook, Schedule, File Watcher | Start workflow |
| **Document** | Read PDF, Write DOCX, OCR | Process files |
| **Transform** | Code, Set, Merge | Manipulate data |
| **Output** | Email, Slack, Google Drive | Deliver results |

### Workflow Example: Contract Review Pipeline

```json
{
  "nodes": [
    {
      "name": "Watch Folder",
      "type": "n8n-nodes-base.localFileTrigger",
      "parameters": {
        "path": "/contracts/incoming",
        "events": ["add"]
      }
    },
    {
      "name": "Extract Text",
      "type": "n8n-nodes-base.readPdf"
    },
    {
      "name": "AI Review",
      "type": "n8n-nodes-base.anthropic",
      "parameters": {
        "model": "claude-sonnet-4-20250514",
        "prompt": "Review this contract for risks..."
      }
    },
    {
      "name": "Save Report",
      "type": "n8n-nodes-base.writeFile"
    },
    {
      "name": "Notify Team",
      "type": "n8n-nodes-base.slack"
    }
  ]
}
```

### Self-Hosting vs Cloud

| Option | Pros | Cons |
|--------|------|------|
| **Self-hosted** | Free, full control, data privacy | Maintenance required |
| **n8n Cloud** | No setup, auto-updates | Costs at scale |

```bash
# Docker quick start
docker run -it --rm \\
  -p 5678:5678 \\
  -v ~/.n8n:/home/node/.n8n \\
  n8nio/n8n
```
""",
        "best_practices": [
            "Start with existing templates, customize as needed",
            "Use error handling nodes for reliability",
            "Store credentials securely with n8n's credential manager",
            "Test workflows with sample data before production"
        ]
    },
    
    "mcp-hub": {
        "name": "mcp-hub",
        "description": "Access 1200+ AI Agent tools via Model Context Protocol (MCP)",
        "tags": ["mcp", "ai-agent", "tools", "integration"],
        "library": {
            "name": "MCP Servers",
            "url": "https://github.com/modelcontextprotocol/servers",
            "stars": "40k+"
        },
        "overview": """This skill provides access to 1200+ MCP (Model Context Protocol) servers - standardized tools that extend AI capabilities. Connect Claude to filesystems, databases, APIs, and document processing tools.""",
        "use_cases": [
            "Access local filesystem to read/write documents",
            "Query databases for data analysis",
            "Integrate with GitHub, Slack, Google Drive",
            "Run document processing tools"
        ],
        "core_concepts": """
### MCP Architecture

```
Claude ←→ MCP Server ←→ External Resource
        (Protocol)      (Files, APIs, DBs)
```

### Popular Document MCP Servers

| Server | Function | Stars |
|--------|----------|-------|
| **filesystem** | Read/write local files | Official |
| **google-drive** | Access Google Docs/Sheets | 5k+ |
| **puppeteer** | Browser automation, PDF gen | 10k+ |
| **sqlite** | Database queries | Official |

### Configuration Example

```json
{
  "mcpServers": {
    "filesystem": {
      "command": "npx",
      "args": [
        "-y",
        "@modelcontextprotocol/server-filesystem",
        "/path/to/documents"
      ]
    },
    "google-drive": {
      "command": "npx",
      "args": ["-y", "@anthropic/mcp-google-drive"]
    }
  }
}
```

### MCP Tool Discovery

Browse available servers:
- [mcp.run](https://mcp.run) - MCP marketplace
- [awesome-mcp-servers](https://github.com/wong2/awesome-mcp-servers)
- [mcp-awesome.com](https://mcp-awesome.com)

### Using MCP in Skills

```python
# MCP tools become available to Claude automatically
# Example: filesystem MCP provides these tools:

# read_file(path) - Read file contents
# write_file(path, content) - Write to file
# list_directory(path) - List directory contents
# search_files(query) - Search for files
```
""",
        "best_practices": [
            "Only enable MCP servers you need (security)",
            "Use official servers when available",
            "Check server permissions before enabling",
            "Combine multiple servers for complex workflows"
        ]
    },
    
    "office-mcp": {
        "name": "office-mcp",
        "description": "MCP server for Word, Excel, PowerPoint operations via AI",
        "tags": ["mcp", "office", "word", "excel", "powerpoint"],
        "library": {
            "name": "Office MCP",
            "url": "https://github.com/anthropics/skills",
            "stars": "N/A"
        },
        "overview": """This skill wraps Office document operations as MCP tools, allowing Claude to create, edit, and manipulate Word, Excel, and PowerPoint files with standardized interfaces.""",
        "use_cases": [
            "Create Word documents with AI-generated content",
            "Build Excel spreadsheets with formulas",
            "Generate PowerPoint presentations",
            "Batch edit Office documents"
        ],
        "core_concepts": """
### Office MCP Tools

| Tool | Input | Output |
|------|-------|--------|
| `create_docx` | Title, sections, styles | .docx file |
| `edit_docx` | Path, changes | Updated .docx |
| `create_xlsx` | Data, formulas | .xlsx file |
| `create_pptx` | Slides, layout | .pptx file |

### Integration with Claude Skills

```markdown
# Example: Combining Skills + MCP

User: "Create a sales report from this data"

1. Data Analysis Skill → Analyze data
2. office-mcp/create_docx → Generate Word report
3. office-mcp/create_xlsx → Generate Excel summary
4. office-mcp/create_pptx → Generate PowerPoint deck
```

### MCP Server Implementation

```python
from mcp import Server
from docx import Document
from openpyxl import Workbook

server = Server("office-mcp")

@server.tool("create_docx")
async def create_docx(title: str, content: str, output_path: str):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    doc.save(output_path)
    return {"status": "success", "path": output_path}

@server.tool("create_xlsx")
async def create_xlsx(data: list, output_path: str):
    wb = Workbook()
    ws = wb.active
    for row in data:
        ws.append(row)
    wb.save(output_path)
    return {"status": "success", "path": output_path}
```
""",
        "best_practices": [
            "Validate inputs before document operations",
            "Use temp files for large documents",
            "Return structured responses with file paths",
            "Handle errors gracefully with meaningful messages"
        ]
    },
    
    "batch-processor": {
        "name": "batch-processor",
        "description": "Process multiple documents in bulk with parallel execution",
        "tags": ["batch", "bulk", "parallel", "automation"],
        "library": {
            "name": "Custom",
            "url": "https://github.com/claude-office-skills/skills",
            "stars": "N/A"
        },
        "overview": """This skill enables efficient bulk processing of documents - convert, transform, extract, or analyze hundreds of files with parallel execution and progress tracking.""",
        "use_cases": [
            "Convert 100 PDFs to Word documents",
            "Extract text from all images in a folder",
            "Batch rename and organize files",
            "Mass update document headers/footers"
        ],
        "core_concepts": """
### Batch Processing Patterns

```
Input: [file1, file2, ..., fileN]
         │
         ▼
    ┌─────────────┐
    │  Parallel   │  ← Process multiple files concurrently
    │  Workers    │
    └─────────────┘
         │
         ▼
Output: [result1, result2, ..., resultN]
```

### Python Implementation

```python
from concurrent.futures import ProcessPoolExecutor, as_completed
from pathlib import Path
from tqdm import tqdm

def process_file(file_path: Path) -> dict:
    \"\"\"Process a single file.\"\"\"
    # Your processing logic here
    return {"path": str(file_path), "status": "success"}

def batch_process(input_dir: str, pattern: str = "*.*", max_workers: int = 4):
    \"\"\"Process all matching files in directory.\"\"\"
    
    files = list(Path(input_dir).glob(pattern))
    results = []
    
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(process_file, f): f for f in files}
        
        for future in tqdm(as_completed(futures), total=len(files)):
            file = futures[future]
            try:
                result = future.result()
                results.append(result)
            except Exception as e:
                results.append({"path": str(file), "error": str(e)})
    
    return results

# Usage
results = batch_process("/documents/invoices", "*.pdf", max_workers=8)
print(f"Processed {len(results)} files")
```

### Error Handling & Resume

```python
import json
from pathlib import Path

class BatchProcessor:
    def __init__(self, checkpoint_file: str = "checkpoint.json"):
        self.checkpoint_file = checkpoint_file
        self.processed = self._load_checkpoint()
    
    def _load_checkpoint(self):
        if Path(self.checkpoint_file).exists():
            return json.load(open(self.checkpoint_file))
        return {}
    
    def _save_checkpoint(self):
        json.dump(self.processed, open(self.checkpoint_file, "w"))
    
    def process(self, files: list, processor_func):
        for file in files:
            if str(file) in self.processed:
                continue  # Skip already processed
            
            try:
                result = processor_func(file)
                self.processed[str(file)] = {"status": "success", **result}
            except Exception as e:
                self.processed[str(file)] = {"status": "error", "error": str(e)}
            
            self._save_checkpoint()  # Resume-safe
```
""",
        "best_practices": [
            "Use progress bars (tqdm) for user feedback",
            "Implement checkpointing for long jobs",
            "Set reasonable worker counts (CPU cores)",
            "Log failures for later review"
        ]
    },
    
    "doc-pipeline": {
        "name": "doc-pipeline",
        "description": "Chain document operations into reusable pipelines",
        "tags": ["pipeline", "workflow", "chain", "automation"],
        "library": {
            "name": "Custom",
            "url": "https://github.com/claude-office-skills/skills",
            "stars": "N/A"
        },
        "overview": """This skill enables building document processing pipelines - chain multiple operations (extract, transform, convert) into reusable workflows with data flowing between stages.""",
        "use_cases": [
            "PDF → Extract Text → Translate → Generate DOCX",
            "Image → OCR → Summarize → Create Report",
            "Excel → Analyze → Generate Charts → Create PPT",
            "Multiple inputs → Merge → Format → Output"
        ],
        "core_concepts": """
### Pipeline Architecture

```
Stage 1      Stage 2      Stage 3      Stage 4
┌──────┐    ┌──────┐    ┌──────┐    ┌──────┐
│Extract│ → │Transform│ → │ AI   │ → │Output│
│ PDF  │    │  Data  │    │Analyze│   │ DOCX │
└──────┘    └──────┘    └──────┘    └──────┘
     │           │           │           │
     └───────────┴───────────┴───────────┘
                 Data Flow
```

### Pipeline DSL (Domain Specific Language)

```yaml
# pipeline.yaml
name: contract-review-pipeline
description: Extract, analyze, and report on contracts

stages:
  - name: extract
    operation: pdf-extraction
    input: $input_file
    output: $extracted_text
    
  - name: analyze
    operation: ai-analyze
    input: $extracted_text
    prompt: "Review this contract for risks..."
    output: $analysis
    
  - name: report
    operation: docx-generation
    input: $analysis
    template: templates/review_report.docx
    output: $output_file
```

### Python Implementation

```python
from typing import Callable, Any
from dataclasses import dataclass

@dataclass
class Stage:
    name: str
    operation: Callable
    
class Pipeline:
    def __init__(self, name: str):
        self.name = name
        self.stages: list[Stage] = []
    
    def add_stage(self, name: str, operation: Callable):
        self.stages.append(Stage(name, operation))
        return self  # Fluent API
    
    def run(self, input_data: Any) -> Any:
        data = input_data
        for stage in self.stages:
            print(f"Running stage: {stage.name}")
            data = stage.operation(data)
        return data

# Example usage
pipeline = Pipeline("contract-review")
pipeline.add_stage("extract", extract_pdf_text)
pipeline.add_stage("analyze", analyze_with_ai)
pipeline.add_stage("generate", create_docx_report)

result = pipeline.run("/path/to/contract.pdf")
```

### Advanced: Conditional Pipelines

```python
class ConditionalPipeline(Pipeline):
    def add_conditional_stage(self, name: str, condition: Callable, 
                               if_true: Callable, if_false: Callable):
        def conditional_op(data):
            if condition(data):
                return if_true(data)
            return if_false(data)
        return self.add_stage(name, conditional_op)

# Usage
pipeline.add_conditional_stage(
    "ocr_if_needed",
    condition=lambda d: d.get("has_images"),
    if_true=run_ocr,
    if_false=lambda d: d
)
```
""",
        "best_practices": [
            "Keep stages focused (single responsibility)",
            "Use intermediate outputs for debugging",
            "Implement stage-level error handling",
            "Make pipelines configurable via YAML/JSON"
        ]
    },
    
    # ========== Template Skills (Missing) ==========
    "invoice-template": {
        "name": "invoice-template",
        "description": "Generate professional PDF invoices from templates",
        "tags": ["invoice", "pdf", "template", "billing"],
        "library": {
            "name": "easy-invoice-pdf",
            "url": "https://github.com/nickmitchko/easy-invoice-pdf",
            "stars": "476"
        },
        "overview": """This skill generates professional PDF invoices from structured data and templates. Create invoices with company branding, itemized lists, tax calculations, and payment details.""",
        "use_cases": [
            "Generate invoices from order data",
            "Create recurring invoices",
            "Batch generate monthly invoices",
            "Customize invoice templates per client"
        ],
        "core_concepts": """
### Invoice Data Structure

```python
invoice_data = {
    "invoice_number": "INV-2026-001",
    "date": "2026-01-30",
    "due_date": "2026-02-28",
    
    "from": {
        "name": "Your Company",
        "address": "123 Business St",
        "email": "billing@company.com"
    },
    
    "to": {
        "name": "Client Name",
        "address": "456 Client Ave",
        "email": "client@example.com"
    },
    
    "items": [
        {"description": "Consulting", "quantity": 10, "rate": 150.00},
        {"description": "Development", "quantity": 20, "rate": 100.00}
    ],
    
    "tax_rate": 0.08,
    "notes": "Payment due within 30 days"
}
```

### PDF Generation with ReportLab

```python
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

def create_invoice(data: dict, output_path: str):
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    
    # Header
    c.setFont("Helvetica-Bold", 24)
    c.drawString(1*inch, height - 1*inch, "INVOICE")
    
    # Invoice details
    c.setFont("Helvetica", 12)
    c.drawString(1*inch, height - 1.5*inch, f"Invoice #: {data['invoice_number']}")
    c.drawString(1*inch, height - 1.75*inch, f"Date: {data['date']}")
    
    # From/To
    y = height - 2.5*inch
    c.drawString(1*inch, y, f"From: {data['from']['name']}")
    c.drawString(4*inch, y, f"To: {data['to']['name']}")
    
    # Items table
    y = height - 4*inch
    c.setFont("Helvetica-Bold", 10)
    c.drawString(1*inch, y, "Description")
    c.drawString(4*inch, y, "Qty")
    c.drawString(5*inch, y, "Rate")
    c.drawString(6*inch, y, "Amount")
    
    c.setFont("Helvetica", 10)
    subtotal = 0
    for item in data['items']:
        y -= 0.3*inch
        amount = item['quantity'] * item['rate']
        subtotal += amount
        c.drawString(1*inch, y, item['description'])
        c.drawString(4*inch, y, str(item['quantity']))
        c.drawString(5*inch, y, f"${item['rate']:.2f}")
        c.drawString(6*inch, y, f"${amount:.2f}")
    
    # Totals
    tax = subtotal * data['tax_rate']
    total = subtotal + tax
    
    y -= 0.5*inch
    c.drawString(5*inch, y, f"Subtotal: ${subtotal:.2f}")
    y -= 0.25*inch
    c.drawString(5*inch, y, f"Tax ({data['tax_rate']*100}%): ${tax:.2f}")
    y -= 0.25*inch
    c.setFont("Helvetica-Bold", 12)
    c.drawString(5*inch, y, f"Total: ${total:.2f}")
    
    c.save()
    return output_path
```

### HTML Template Approach

```python
from weasyprint import HTML
from jinja2 import Template

invoice_template = \"\"\"
<!DOCTYPE html>
<html>
<head>
    <style>
        body { font-family: Arial; margin: 40px; }
        .header { display: flex; justify-content: space-between; }
        table { width: 100%; border-collapse: collapse; margin: 20px 0; }
        th, td { border: 1px solid #ddd; padding: 10px; text-align: left; }
        .total { font-weight: bold; font-size: 18px; }
    </style>
</head>
<body>
    <div class="header">
        <h1>INVOICE</h1>
        <div>
            <p>Invoice #: {{ invoice_number }}</p>
            <p>Date: {{ date }}</p>
        </div>
    </div>
    <table>
        <tr><th>Description</th><th>Qty</th><th>Rate</th><th>Amount</th></tr>
        {% for item in items %}
        <tr>
            <td>{{ item.description }}</td>
            <td>{{ item.quantity }}</td>
            <td>${{ "%.2f"|format(item.rate) }}</td>
            <td>${{ "%.2f"|format(item.quantity * item.rate) }}</td>
        </tr>
        {% endfor %}
    </table>
    <p class="total">Total: ${{ "%.2f"|format(total) }}</p>
</body>
</html>
\"\"\"

def create_invoice_html(data: dict, output_path: str):
    template = Template(invoice_template)
    
    # Calculate total
    total = sum(i['quantity'] * i['rate'] for i in data['items'])
    total *= (1 + data.get('tax_rate', 0))
    data['total'] = total
    
    html = template.render(**data)
    HTML(string=html).write_pdf(output_path)
    return output_path
```
""",
        "best_practices": [
            "Validate required fields before generation",
            "Use templates for consistent branding",
            "Auto-calculate totals (don't trust input)",
            "Include payment instructions and terms"
        ]
    },
    
    "template-engine": {
        "name": "template-engine",
        "description": "Auto-fill document templates with data - mail merge for any format",
        "tags": ["template", "mail-merge", "autofill", "automation"],
        "library": {
            "name": "docxtpl / yumdocs",
            "url": "https://github.com/elapouya/python-docxtpl",
            "stars": "2.1k"
        },
        "overview": """This skill enables template-based document generation - define templates with placeholders, then automatically fill them with data. Works with Word, Excel, PowerPoint, and more.""",
        "use_cases": [
            "Mail merge for bulk letters/contracts",
            "Generate personalized reports from data",
            "Create certificates from templates",
            "Auto-fill forms with user data"
        ],
        "core_concepts": """
### Template Syntax (Jinja2-based)

```
{{ variable }}           - Simple substitution
{% for item in list %}   - Loop
{% if condition %}       - Conditional
{{ date | format_date }} - Filter
```

### Word Template Example

```python
from docxtpl import DocxTemplate

# Create template with placeholders:
# Dear {{ name }},
# Thank you for your order #{{ order_id }}...

def fill_template(template_path: str, data: dict, output_path: str):
    doc = DocxTemplate(template_path)
    doc.render(data)
    doc.save(output_path)
    return output_path

# Usage
fill_template(
    "templates/order_confirmation.docx",
    {
        "name": "John Smith",
        "order_id": "ORD-12345",
        "items": [
            {"name": "Product A", "qty": 2, "price": 29.99},
            {"name": "Product B", "qty": 1, "price": 49.99}
        ],
        "total": 109.97
    },
    "output/confirmation_john.docx"
)
```

### Excel Template

```python
from openpyxl import load_workbook
import re

def fill_excel_template(template_path: str, data: dict, output_path: str):
    wb = load_workbook(template_path)
    ws = wb.active
    
    # Find and replace placeholders like {{name}}
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                for key, value in data.items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in cell.value:
                        cell.value = cell.value.replace(placeholder, str(value))
    
    wb.save(output_path)
    return output_path
```

### Bulk Generation (Mail Merge)

```python
import csv
from pathlib import Path

def mail_merge(template_path: str, data_csv: str, output_dir: str):
    \"\"\"Generate documents for each row in CSV.\"\"\"
    
    Path(output_dir).mkdir(exist_ok=True)
    
    with open(data_csv) as f:
        reader = csv.DictReader(f)
        
        for i, row in enumerate(reader):
            output_path = f"{output_dir}/document_{i+1}.docx"
            fill_template(template_path, row, output_path)
            print(f"Generated: {output_path}")

# Usage with contacts.csv:
# name,email,company
# John,john@example.com,Acme
# Jane,jane@example.com,Corp

mail_merge(
    "templates/welcome_letter.docx",
    "data/contacts.csv",
    "output/letters"
)
```

### Advanced: Conditional Content

```python
from docxtpl import DocxTemplate

# Template with conditionals:
# {% if vip %}
# Thank you for being a VIP member!
# {% else %}
# Thank you for your purchase.
# {% endif %}

doc = DocxTemplate("template.docx")
doc.render({
    "name": "John",
    "vip": True,
    "discount": 20
})
doc.save("output.docx")
```
""",
        "best_practices": [
            "Use clear placeholder naming ({{client_name}})",
            "Validate data before rendering",
            "Handle missing data gracefully",
            "Keep templates version-controlled"
        ]
    }
}

def create_skill_md(skill_config: dict) -> str:
    """Generate SKILL.md content from config."""
    
    content = f"""---
name: {skill_config['name']}
description: {skill_config['description']}
author: claude-office-skills
version: "1.0"
tags: {skill_config['tags']}
models: [claude-sonnet-4, claude-opus-4]
tools: [computer, code_execution, file_operations]
library:
  name: {skill_config['library']['name']}
  url: {skill_config['library']['url']}
  stars: {skill_config['library']['stars']}
---

# {skill_config['name'].replace('-', ' ').title()} Skill

## Overview

{skill_config['overview']}

## How to Use

1. Describe what you want to accomplish
2. Provide any required input data or files
3. I'll execute the appropriate operations

**Example prompts:**
"""
    
    for use_case in skill_config['use_cases']:
        content += f'- "{use_case}"\n'
    
    content += f"""
## Domain Knowledge

{skill_config['core_concepts']}

## Best Practices

"""
    
    for i, practice in enumerate(skill_config['best_practices'], 1):
        content += f"{i}. **{practice.split(':')[0] if ':' in practice else practice}**\n"
    
    content += f"""
## Installation

```bash
# Install required dependencies
pip install python-docx openpyxl python-pptx reportlab jinja2
```

## Resources

- [{skill_config['library']['name']} Repository]({skill_config['library']['url']})
- [Claude Office Skills Hub](https://github.com/claude-office-skills/skills)
"""
    
    return content

def main():
    """Create all missing skills."""
    
    print("=" * 60)
    print("Claude Office Skills - Batch Creator")
    print(f"Creating {len(SKILLS_TO_CREATE)} skills...")
    print("=" * 60)
    
    created = []
    skipped = []
    
    for skill_name, config in SKILLS_TO_CREATE.items():
        skill_dir = BASE_DIR / skill_name
        skill_file = skill_dir / "SKILL.md"
        
        if skill_file.exists():
            print(f"⏭️  Skipping {skill_name} (already exists)")
            skipped.append(skill_name)
            continue
        
        # Create directory
        skill_dir.mkdir(exist_ok=True)
        
        # Generate and write SKILL.md
        content = create_skill_md(config)
        skill_file.write_text(content)
        
        print(f"✅ Created {skill_name}/SKILL.md")
        created.append(skill_name)
    
    print("\n" + "=" * 60)
    print(f"Summary:")
    print(f"  Created: {len(created)} skills")
    print(f"  Skipped: {len(skipped)} skills (already exist)")
    print("=" * 60)
    
    if created:
        print("\nNew skills created:")
        for skill in created:
            print(f"  - {skill}/SKILL.md")
    
    return created

if __name__ == "__main__":
    main()
